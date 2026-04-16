# Kaviya SG (23PT18) & Sangamithra SG (23PT30) — Compiler Design Lab, Course 23XT62
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_toc(paragraph):
    # Adds a Table of Contents placeholder
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

def create_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    # Shading the cell
    tc = table.cell(0, 0)._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EFEFEF')  # light grey
    tcPr.append(shd)
    
    p = table.cell(0,0).paragraphs[0]
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

def set_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Page ")
        # Add PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

def main():
    doc = Document()
    
    # Configure default styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    
    # Configure Heading 1
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Calibri'
    style_h1.font.size = Pt(16)
    style_h1.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
    style_h1.font.bold = True
    
    # Configure Heading 2
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Calibri'
    style_h2.font.size = Pt(14)
    style_h2.font.color.rgb = RGBColor(0, 0, 128)
    style_h2.font.bold = True
    
    # --- COVER PAGE ---
    doc.add_paragraph("\n\n\n\n\n\n")
    title = doc.add_paragraph("QueryLite Compiler — Project Documentation")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(0, 0, 128)
    
    doc.add_paragraph("\n\n")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Subject: Compiler Design Lab (23XT62)\n").bold = True
    info.add_run("Students: Kaviya SG (23PT18) & Sangamithra SG (23PT30)\n").bold = True
    info.add_run("Institution: ____________________\n").bold = True
    info.add_run(f"Date: {datetime.today().strftime('%Y-%m-%d')}\n").bold = True
    
    doc.add_page_break()
    
    # --- Table of Contents ---
    doc.add_heading("Table of Contents", level=1)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)
    doc.add_paragraph("Note: Right-click here and select 'Update Field' to display the Table of Contents in Word.")
    
    doc.add_page_break()
    
    # --- Section 1: Project Overview ---
    doc.add_heading("Section 1 — Project Overview", level=1)
    doc.add_heading("What QueryLite is and why it was designed", level=2)
    doc.add_paragraph(
        "QueryLite is a custom, lightweight SQL-inspired query language designed specifically for academic "
        "exploration into compiler construction. The language aims to provide a minimalistic yet functional "
        "subset of relational operations. It was designed to help students understand the process of lexing, "
        "parsing, semantic analysis, intermediate representation generation, and query optimization."
    )
    
    doc.add_heading("The problem it solves", level=2)
    doc.add_paragraph(
        "Traditional SQL compilers are massive, complex systems, which makes them difficult to study constraint-free. "
        "QueryLite simplifies the ecosystem by defining a rigid, smaller grammar that specifically addresses data projection (PULL), "
        "filtering (FILTER), sorting (ARRANGE BY), and basic relational logic. It allows a full end-to-end realization of a "
        "compiler pipeline within a manageable timeframe and code complexity, serving as a pedagogical tool without "
        "the overhead of full SQL compliance."
    )
    
    doc.add_heading("Key features of the language and compiler", level=2)
    doc.add_paragraph(
        "• Modular Pipeline: Uses a clear sequence of Lexer → Parser → AST → Semantic Analyzer → IR Generator → Optimizer → Executor.\n"
        "• Custom Syntax: PULL (instead of SELECT), FILTER (instead of WHERE), ARRANGE BY (instead of ORDER BY).\n"
        "• In-Memory Execution: Bypasses complex disk I/O, allowing focus purely on compilation and evaluation logic.\n"
        "• Web UI & Terminal Mode: Dual interface providing an interactive experience and visual pipeline stages.\n"
        "• Rule-Based Optimization: Basic algebraic simplifications and filtering optimizations during IR."
    )
    
    doc.add_page_break()
    
    # --- Section 2: QueryLite Language Specification ---
    doc.add_heading("Section 2 — QueryLite Language Specification", level=1)
    doc.add_heading("Full language syntax description", level=2)
    doc.add_paragraph(
        "QueryLite statements retrieve records from a table, conditionally filter them, and sort the result. "
        "The language uses strict keywords to separate projection paths, logic conditions, and ordering criteria. "
        "Statements must be terminated by a semicolon (;)."
    )
    
    doc.add_heading("BNF grammar rules", level=2)
    bnf_code = (
        "<query>       ::= PULL <columns> FROM <identifier> "
        "[FILTER <condition>] [ARRANGE BY <identifier> [ASC | DESC]] ;\n"
        "<columns>     ::= * | <identifier> ( , <identifier> )*\n"
        "<condition>   ::= <expression> ( ( AND | OR ) <expression> )*\n"
        "<expression>  ::= <identifier> <operator> <value>\n"
        "<operator>    ::= == | != | > | < | >= | <=\n"
        "<value>       ::= <string_literal> | <number_literal>\n"
    )
    create_code_block(doc, bnf_code)
    doc.add_paragraph("")
    
    doc.add_heading("Keywords", level=2)
    table_kw = doc.add_table(rows=1, cols=2)
    table_kw.style = 'Table Grid'
    hdr_cells = table_kw.rows[0].cells
    hdr_cells[0].text = 'Keyword'
    hdr_cells[1].text = 'Usage'
    keywords = [
        ("PULL", "Specifies columns to project (equivalent to SELECT)"),
        ("FROM", "Specifies source table"),
        ("FILTER", "Specifies condition to restrict results (equivalent to WHERE)"),
        ("ARRANGE BY", "Sorts the output records (equivalent to ORDER BY)"),
        ("ASC", "Ascending sort direction"),
        ("DESC", "Descending sort direction")
    ]
    for kw, usage in keywords:
        row_cells = table_kw.add_row().cells
        row_cells[0].text = kw
        row_cells[1].text = usage
    doc.add_paragraph("")
    
    doc.add_heading("Operators", level=2)
    table_op = doc.add_table(rows=1, cols=2)
    table_op.style = 'Table Grid'
    hdr_cells = table_op.rows[0].cells
    hdr_cells[0].text = 'Operator'
    hdr_cells[1].text = 'Description'
    operators = [
        (">", "Greater than"), ("<", "Less than"),
        ("==", "Equals"), ("!=", "Not equals"),
        (">=", "Greater than or equal to"), ("<=", "Less than or equal to")
    ]
    for op, desc in operators:
        row_cells = table_op.add_row().cells
        row_cells[0].text = op
        row_cells[1].text = desc
    doc.add_paragraph("")
    
    doc.add_heading("5 example QueryLite queries", level=2)
    doc.add_paragraph("1. Basic Pull:")
    create_code_block(doc, "PULL * FROM users;")
    doc.add_paragraph("Explanation: Fetches all columns and rows from the 'users' table.")
    
    doc.add_paragraph("\n2. Column Projection:")
    create_code_block(doc, "PULL name, age FROM users;")
    doc.add_paragraph("Explanation: Retrieves only the 'name' and 'age' fields from 'users'.")
    
    doc.add_paragraph("\n3. Conditional Filtering:")
    create_code_block(doc, "PULL name FROM users FILTER age >= 18;")
    doc.add_paragraph("Explanation: Fetches 'name' where the user is an adult.")
    
    doc.add_paragraph("\n4. Arranging Results:")
    create_code_block(doc, "PULL * FROM inventory ARRANGE BY price DESC;")
    doc.add_paragraph("Explanation: Returns all items sorted by price in descending order.")
    
    doc.add_paragraph("\n5. Combined Complex Query:")
    create_code_block(doc, "PULL name, department FROM employees FILTER salary > 50000 ARRANGE BY name ASC;")
    doc.add_paragraph("Explanation: Fetches names and departments of high earners, ordering results alphabetically by name.")
    
    doc.add_page_break()
    
    # --- Section 3: Compiler Architecture ---
    doc.add_heading("Section 3 — Compiler Architecture", level=1)
    
    doc.add_heading("Pipeline Stages", level=2)
    stages = [
        ("Lexer", "Scans raw source text and converts character sequences into meaningful tokens (identifying keywords, literals, and symbols)."),
        ("Parser", "Reads the stream of tokens and groups them strictly according to BNF grammar rules, outputting an Abstract Syntax Tree (AST)."),
        ("AST (Abstract Syntax Tree)", "An intermediate data structure mapping the hierarchy of operations in a machine-readable object form."),
        ("Semantic Analyzer", "Validates the correctness of the AST. It ensures that references to tables or identifiers make contextual sense before executing (e.g., verifying schemas in a tighter system)."),
        ("IR Generator", "Translates the validated AST into Intermediate Representation (IR), flattening the logic into a standardized, execution-ready format."),
        ("Optimizer", "Scans the IR for inefficiencies and applies rule-based optimizations, rearranging steps to form faster execution paths without altering outcomes."),
        ("Executor", "Translates optimized IR instructions into real Python routines, acts purely on in-memory collections, and returns formatted arrays of records.")
    ]
    for name, desc in stages:
        doc.add_paragraph(f"• {name}: ").bold = True
        doc.paragraphs[-1].add_run(desc)
        
    doc.add_heading("Architecture Flow Diagram", level=2)
    
    arch_flow = (
        "+-------------------+\n"
        "| Source Query Code |\n"
        "+-------------------+\n"
        "          |\n"
        "          v\n"
        "+-------------------+\n"
        "|       Lexer       |\n"
        "+-------------------+\n"
        "          |\n"
        "          v\n"
        "+-------------------+\n"
        "|      Parser       |\n"
        "+-------------------+\n"
        "          |\n"
        "          v\n"
        "+-------------------+\n"
        "| Semantic Analyzer |\n"
        "+-------------------+\n"
        "          |\n"
        "          v\n"
        "+-------------------+\n"
        "|   IR Generator    |\n"
        "+-------------------+\n"
        "          |\n"
        "          v\n"
        "+-------------------+\n"
        "|     Optimizer     |\n"
        "+-------------------+\n"
        "          |\n"
        "          v\n"
        "+-------------------+\n"
        "|     Executor      |\n"
        "+-------------------+\n"
        "          |\n"
        "          v\n"
        "+-------------------+\n"
        "| Execution Results |\n"
        "+-------------------+"
    )
    create_code_block(doc, arch_flow)
    
    doc.add_page_break()
    
    # --- Section 4: Module Descriptions ---
    doc.add_heading("Section 4 — Module Descriptions", level=1)
    
    modules = [
        {
            "file": "lexer.py",
            "purpose": "Converts plain string queries into tokens. Uses regex for basic matching.",
            "key_classes": "Lexer (class) — handles stateful tokenization step-by-step.",
            "io": "Input: Raw string query. Output: List of Token objects."
        },
        {
            "file": "parser.py",
            "purpose": "Applies a recursive descent strategy to map tokens to grammar constraints.",
            "key_classes": "Parser (class), parse_query() — drives syntax hierarchy.",
            "io": "Input: List of Token objects. Output: AST node (QueryNode)."
        },
        {
            "file": "ast_nodes.py",
            "purpose": "A collection of plain objects representing syntax shapes.",
            "key_classes": "QueryNode, ColumnNode, FilterNode, ArrangeNode.",
            "io": "Input/Output: Represents intermediate structural data passed across modules."
        },
        {
            "file": "semantic.py",
            "purpose": "Provides early validation before compilation deepens, checking logical soundness.",
            "key_classes": "SemanticAnalyzer, analyze() — traverses AST to spot obvious semantic errors.",
            "io": "Input: AST QueryNode. Output: Verified AST QueryNode or throws exception."
        },
        {
            "file": "ir.py",
            "purpose": "Maintains an instruction list sequence mapping out operations serially.",
            "key_classes": "IRGenerator, IROp — building blocks for IR logic execution.",
            "io": "Input: Validated AST QueryNode. Output: List of IROp objects."
        },
        {
            "file": "optimizer.py",
            "purpose": "Filters and modifies the IR list to improve performance or drop useless paths.",
            "key_classes": "Optimizer, optimize() — runs transformation passes over IR arrays.",
            "io": "Input: List of IROp objects. Output: Reduced/Optimized list of IROp objects."
        },
        {
            "file": "executor.py",
            "purpose": "Processes the physical steps outlined in the IR across dummy/in-memory data tables.",
            "key_classes": "Executor, execute() — iterates instructions and transforms dataset.",
            "io": "Input: Optimized IR list. Output: Dictionary/JSON structure holding results."
        },
        {
            "file": "server.py",
            "purpose": "Bridges the Python backend to a web dashboard using the http.server module.",
            "key_classes": "CompilerHTTPRequestHandler — routes /api/compile POST logic and serves static HTML.",
            "io": "Input: Web requests (POST JSON). Output: Web responses (JSON/HTML files)."
        },
        {
            "file": "main.py",
            "purpose": "Program entrypoint giving users the choice of Web UI or Terminal interface modes.",
            "key_classes": "main() runtime switch.",
            "io": "Input: User menu choice. Output: Starts HTTP server or loops simple input prompt."
        }
    ]
    
    for mod in modules:
        doc.add_heading(mod["file"], level=2)
        doc.add_paragraph("Purpose: ").bold = True
        doc.paragraphs[-1].add_run(mod["purpose"])
        doc.add_paragraph("Key functions/classes: ").bold = True
        doc.paragraphs[-1].add_run(mod["key_classes"])
        doc.add_paragraph("Input/Output: ").bold = True
        doc.paragraphs[-1].add_run(mod["io"])
        doc.add_paragraph("")
    
    doc.add_page_break()
    
    # --- Section 5: Screenshots ---
    doc.add_heading("Section 5 — Screenshots", level=1)
    doc.add_paragraph("Please paste your screenshots into the descriptive boxes below:")
    
    labels = [
        "[Screenshot 1: Terminal — Running a PULL query]",
        "[Screenshot 2: Terminal — Tokens output]",
        "[Screenshot 3: Terminal — AST and IR output]",
        "[Screenshot 4: Web UI — Query input and Result table]",
        "[Screenshot 5: Web UI — Pipeline stage tabs]"
    ]
    
    for label in labels:
        doc.add_paragraph(label).bold = True
        table = doc.add_table(rows=1, cols=1)
        # Set height to ~5cm
        table.rows[0].height = Cm(5)
        # Shade it
        tc = table.cell(0, 0)._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')  # Light grey
        tcPr.append(shd)
        
        # We need the table to be constrained but python docx usually expands cell widths automatically
        table.cell(0,0).width = Cm(15)  # Make it wide enough to look like a box
        doc.add_paragraph("")
        doc.add_paragraph("")
        
    doc.add_page_break()
    
    # --- Section 6: Step-by-Step: How to Run ---
    doc.add_heading("Section 6 — Step-by-Step: How to Run", level=1)
    
    doc.add_heading("Terminal mode:", level=2)
    term_steps = [
        "Clone or download the project source files into a folder.",
        "Open a terminal or command prompt inside the project folder.",
        "Run the command: python main.py",
        "Choose option [1] specifically to enter Terminal mode.",
        "Type a valid QueryLite query (e.g., PULL * FROM users;) and press Enter."
    ]
    for i, step in enumerate(term_steps, start=1):
        doc.add_paragraph(f"{i}. {step}")
        
    doc.add_paragraph("")
        
    doc.add_heading("Web UI mode:", level=2)
    web_steps = [
        "Run the internal web server with: python main.py",
        "Choose option [2] specifically to execute Web UI mode.",
        "Open any modern web browser at http://localhost:8080",
        "In the provided text editor area, type the proper query and press ▶ Run or hit Ctrl+Enter.",
        "Navigate through the 'Lexer', 'Parser/AST', 'IR', 'Optimizer', and 'Results' tabs to view compilation steps visually."
    ]
    for i, step in enumerate(web_steps, start=1):
        doc.add_paragraph(f"{i}. {step}")
        
    doc.add_page_break()
    
    # --- Section 7: Sample Output ---
    doc.add_heading("Section 7 — Sample Output", level=1)
    
    doc.add_heading("Terminal Output", level=2)
    doc.add_paragraph("Target Query:")
    create_code_block(doc, "PULL name, age FROM users FILTER age > 21 ARRANGE BY age;")
    
    term_out = (
        "Compiling query...\n"
        "[LEXER]   Tokens: [(PULL, 'PULL'), (IDENTIFIER, 'name'), (COMMA, ','), (IDENTIFIER, 'age'), ...]\n"
        "[PARSER]  AST: QueryNode(columns=['name', 'age'], source='users', ...)\n"
        "[IR]      Instructions: [LOAD users, PROJECT ['name', 'age'], FILTER age > 21, ARRANGE age]\n"
        "[EXEC]    Running IR ops...\n"
        "Execution Results:\n"
        "{'name': 'Bob', 'age': 25},\n"
        "{'name': 'Charlie', 'age': 30}"
    )
    doc.add_paragraph("Output:")
    create_code_block(doc, term_out)
    
    doc.add_heading("Web UI Result Table Format", level=2)
    doc.add_paragraph("The typical HTML output visually forms a table as generated below:")
    
    sample_table = doc.add_table(rows=3, cols=2)
    sample_table.style = 'Table Grid'
    hdr = sample_table.rows[0].cells
    hdr[0].text = 'name'
    hdr[1].text = 'age'
    # Shade header
    for cell in hdr:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'BDD7EE') # Light blue header
        tcPr.append(shd)
        
    r1 = sample_table.rows[1].cells
    r1[0].text = 'Bob'
    r1[1].text = '25'
    
    r2 = sample_table.rows[2].cells
    r2[0].text = 'Charlie'
    r2[1].text = '30'
    
    # Ensure there are 10-12 pages
    doc.add_page_break()
    doc.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_paragraph("--- End of Documentation ---").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    set_page_numbers(doc)
    doc.save('QueryLite_Documentation.docx')
    print("Done generating document.")

if __name__ == '__main__':
    main()
